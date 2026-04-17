#!/usr/bin/env python3
"""
NetSuite PCI File Scanner

Validates flagged files for PCI-DSS compliance by:
1. Parsing a flagged file list (pipe-delimited, Windows paths)
2. Resolving NetSuite file IDs via SuiteQL (subfolder name + filename matching)
3. Downloading each file from NetSuite
4. Extracting text (PDF, Excel, Word, OCR fallback for scanned PDFs)
5. Scanning for credit card patterns (regex + Luhn checksum)
6. Generating a JSON report + human-readable summary

Usage:
  python3 pci_scan.py --file-list /tmp/confirmed_files.txt
  python3 pci_scan.py --file-list /tmp/confirmed_files.txt --output-dir /tmp/pci-results
  python3 pci_scan.py --file-list /tmp/confirmed_files.txt --resolve-only
  python3 pci_scan.py --file-list /tmp/confirmed_files.txt --skip-download --scan-dir /tmp/existing-downloads
"""

import sys
import os
import re
import json
import argparse
import base64
import time
import urllib.request
import urllib.error
from pathlib import Path
from datetime import datetime
from typing import Dict, Any, Optional, List, Tuple, Set

# ---------------------------------------------------------------------------
# NetSuite API Gateway (same pattern as query_netsuite.py)
# ---------------------------------------------------------------------------

_gw_base = os.environ.get('NETSUITE_GATEWAY_URL', 'https://nsapi.twistedx.tech').rstrip('/')
GATEWAY_URL = f'{_gw_base}/api/suiteapi'
_API_KEY = os.environ.get('NETSUITE_API_KEY', '')

DEFAULT_ACCOUNT = 'twistedx'
DEFAULT_ENVIRONMENT = 'production'
ROOT_FOLDER_ID = 18625  # "Twisted X Attachments" root
RATE_LIMIT_DELAY = 0.2  # seconds between downloads
QUERY_BATCH_SIZE = 50   # max items per IN clause

ACCOUNT_ALIASES = {
    'twx': 'twistedx', 'twisted': 'twistedx', 'twistedx': 'twistedx',
}
ENV_ALIASES = {
    'prod': 'production', 'production': 'production',
    'sb1': 'sandbox', 'sandbox': 'sandbox',
}


def _gateway_headers() -> dict:
    """Return headers for API gateway requests, including API key if set."""
    if _API_KEY:
        return {'Content-Type': 'application/json', 'Accept': 'application/json', 'X-API-Key': _API_KEY}
    return {'Content-Type': 'application/json', 'Accept': 'application/json', 'Origin': _gw_base}


def resolve_account(account: str) -> str:
    return ACCOUNT_ALIASES.get(account.lower(), account.lower())


def resolve_environment(environment: str) -> str:
    return ENV_ALIASES.get(environment.lower(), environment.lower())


def execute_query(
    query: str,
    params: Optional[List[Any]] = None,
    account: str = DEFAULT_ACCOUNT,
    environment: str = DEFAULT_ENVIRONMENT,
    timeout: int = 300
) -> Dict[str, Any]:
    """Execute a SuiteQL query via the API Gateway."""
    payload = {
        'action': 'queryRun',
        'procedure': 'queryRun',
        'query': query,
        'params': params or [],
        'returnAllRows': True,
        'netsuiteAccount': resolve_account(account),
        'netsuiteEnvironment': resolve_environment(environment),
    }
    try:
        data = json.dumps(payload).encode('utf-8')
        req = urllib.request.Request(GATEWAY_URL, data=data, headers=_gateway_headers())
        with urllib.request.urlopen(req, timeout=timeout) as resp:
            result = json.loads(resp.read().decode('utf-8'))
            records = (
                result.get('data', {}).get('records', [])
                or result.get('records', [])
            )
            return {'records': records, 'error': None}
    except urllib.error.HTTPError as e:
        return {'records': [], 'error': f"HTTP {e.code}: {e.read().decode('utf-8') if e.fp else str(e)}"}
    except Exception as e:
        return {'records': [], 'error': str(e)}


def _decode_gateway_content(content_b64: str) -> bytes:
    """
    Decode file content from the NetSuite API Gateway.
    The gateway double-encodes binary files:
      RESTlet returns base64(raw_bytes) → gateway wraps → base64(base64(raw_bytes))
    First decode produces an inner base64 string; second decode gives raw bytes.
    """
    first_decoded = base64.b64decode(content_b64)
    try:
        intermediate = first_decoded.decode('utf-8')
        return base64.b64decode(intermediate)
    except Exception:
        return first_decoded


def download_file_content(
    file_id: int,
    account: str = DEFAULT_ACCOUNT,
    environment: str = DEFAULT_ENVIRONMENT,
) -> Optional[bytes]:
    """Download file bytes from NetSuite via fileGet."""
    payload = {
        'action': 'fileGet',
        'procedure': 'fileGet',
        'id': file_id,
        'returnContent': True,
        'netsuiteAccount': resolve_account(account),
        'netsuiteEnvironment': resolve_environment(environment),
    }
    try:
        data = json.dumps(payload).encode('utf-8')
        req = urllib.request.Request(GATEWAY_URL, data=data, headers=_gateway_headers())
        with urllib.request.urlopen(req, timeout=120) as resp:
            result = json.loads(resp.read().decode('utf-8'))
            if not result.get('success'):
                return None
            d = result.get('data', {})
            # Format B: data.file.content (double-base64 encoded)
            if isinstance(d, dict) and 'file' in d:
                content = d['file'].get('content')
                if content:
                    return _decode_gateway_content(content)
            # Format A: data.content
            if isinstance(d, dict) and d.get('content'):
                content = d['content']
                if d.get('encoding') == 'BASE64' or d.get('isBase64'):
                    return _decode_gateway_content(content)
                return content.encode('utf-8')
            # Fallback
            if isinstance(result.get('file'), dict) and result['file'].get('content'):
                return _decode_gateway_content(result['file']['content'])
            return None
    except Exception:
        return None


# ---------------------------------------------------------------------------
# Flagged file list parsing
# ---------------------------------------------------------------------------

def parse_flagged_list(file_path: str) -> List[Dict[str, str]]:
    """
    Parse pipe-delimited flagged file list.
    Format: ?|?|filename|windows_path
    Extract subfolder name and filename from the Windows path.
    """
    entries = []
    with open(file_path, encoding='utf-8', errors='replace') as f:
        for lineno, line in enumerate(f, 1):
            line = line.strip()
            if not line:
                continue
            parts = line.split('|')
            if len(parts) < 4:
                print(f"  WARN: line {lineno} has fewer than 4 fields, skipping: {line[:80]}", file=sys.stderr)
                continue
            win_path = parts[3].replace('\\', '/')
            path_parts = [p for p in win_path.split('/') if p]
            if len(path_parts) < 2:
                print(f"  WARN: line {lineno} path too short: {win_path}", file=sys.stderr)
                continue
            subfolder_name = path_parts[-2]
            filename = path_parts[-1]
            entries.append({
                'lineno': lineno,
                'subfolder_name': subfolder_name,
                'filename': filename,
                'original_line': line,
            })
    return entries


# ---------------------------------------------------------------------------
# Filename dedup normalization (Windows _N_ -> NetSuite (N))
# ---------------------------------------------------------------------------

def normalize_filename(win_filename: str) -> List[str]:
    """
    Generate candidate NetSuite filenames from a Windows backup filename.
    The Windows scanner sanitized special characters to underscores:
      (N)          -> _N_   (dedup numbers)
      [N]          -> _N_   (square bracket numbers)
      [TEXT]       -> _TEXT_ (square bracket text)
      (TEXT)       -> _TEXT_ (parenthesized text)
      #            -> _      (hash/pound)
      $            -> _      (dollar sign, leading)
      '            -> _      (apostrophe)
      ,            -> _      (comma in numbers)
    Returns candidates in priority order (most specific first).
    """
    candidates = [win_filename]  # exact match first

    # --- Dedup suffix patterns (space-prefixed _N_) ---
    # " _N_" -> " (N)"
    ns_paren = re.sub(r' _(\d+)_', r' (\1)', win_filename)
    _add_unique(candidates, ns_paren)

    # " _N_" -> " [N]"
    ns_bracket = re.sub(r' _(\d+)_', r' [\1]', win_filename)
    _add_unique(candidates, ns_bracket)

    # "_N_" without leading space -> "[N]" (e.g. Letter_6_.docx, Authorization_49369_.docx)
    ns_bracket2 = re.sub(r'_(\d+)_', r'[\1]', win_filename)
    _add_unique(candidates, ns_bracket2)

    # "_N_" -> "(N)"
    ns_paren2 = re.sub(r'_(\d+)_', r'(\1)', win_filename)
    _add_unique(candidates, ns_paren2)

    # --- Underscore as special character ---
    # Trailing lone "_" before extension or at end -> "#"
    # e.g. "CORRECT FAX _.pdf" -> "CORRECT FAX #.pdf"
    ns_hash1 = re.sub(r' _(?=\.|\s|$)', ' #', win_filename)
    _add_unique(candidates, ns_hash1)

    # "_" in word context (not adjacent to digits) -> "#"
    # e.g. "Transaction _490897..." -> "Transaction #490897..."
    ns_hash2 = re.sub(r' _(\d)', r' #\1', win_filename)
    _add_unique(candidates, ns_hash2)

    # Apply hash sub to dedup-stripped variants too
    if ns_paren != win_filename:
        _add_unique(candidates, re.sub(r' _(?=\.|\s|$)', ' #', ns_paren))
        _add_unique(candidates, re.sub(r' _(\d)', r' #\1', ns_paren))

    # Leading "_" before digits/letters -> "$"
    # e.g. "_1726.27.pdf" -> "$1726.27.pdf"
    if win_filename.startswith('_'):
        _add_unique(candidates, '$' + win_filename[1:])
        # Also handle "_N_545.83.pdf" -> "$N,545.83.pdf" (comma in number)
        dollar_ver = '$' + win_filename[1:]
        comma_ver = re.sub(r'_(\d)', r',\1', dollar_ver)
        _add_unique(candidates, comma_ver)

    # "_s " -> "'s " (apostrophe-s)
    ns_apos = re.sub(r"_s\b", "'s", win_filename)
    _add_unique(candidates, ns_apos)
    if ns_apos != win_filename:
        # Also apply hash transform to the apostrophe version
        _add_unique(candidates, re.sub(r' _(\d)', r' #\1', ns_apos))

    # --- Strip all dedup suffixes to get bare base name ---
    base = re.sub(r'(\s*_\d+_)+(?=\.\w+$)', '', win_filename)
    _add_unique(candidates, base)

    base_bracket = re.sub(r'(\s*\[\d+\])+(?=\.\w+$)', '', ns_bracket2)
    _add_unique(candidates, base_bracket)

    base_paren = re.sub(r'(\s*\(\d+\))+(?=\.\w+$)', '', ns_paren)
    _add_unique(candidates, base_paren)

    return candidates


def _add_unique(lst: List[str], item: str) -> None:
    """Append item to list only if not already present."""
    if item not in lst:
        lst.append(item)


def fuzzy_normalize(name: str) -> str:
    """Strip extension and all non-alphanumeric chars for fuzzy comparison."""
    stem = Path(name).stem.lower()
    return re.sub(r'[^a-z0-9]', '', stem)


# ---------------------------------------------------------------------------
# NetSuite ID resolution
# ---------------------------------------------------------------------------

def _win_name_to_ns_candidates(win_name: str) -> List[str]:
    """
    Convert a Windows-sanitized folder name back to NetSuite candidates.
    Windows replaces {  } with _ _ and curly braces with underscores.
    e.g.  _temp-2025-11-26-10-31-650-36_  ->  {temp-2025-11-26-10-31-650-36}
    """
    candidates = [win_name]
    # _temp-..._  ->  {temp-...}
    ns_curly = re.sub(r'^_(temp-.+)_$', r'{\1}', win_name)
    if ns_curly != win_name:
        candidates.append(ns_curly)
    return candidates


def resolve_folder_ids(
    subfolder_names: Set[str],
    account: str,
    environment: str,
) -> Dict[str, int]:
    """
    Map subfolder names -> NetSuite folder internal IDs.
    First searches under ROOT_FOLDER_ID (18625) tree.
    For any unresolved names, falls back to a broader org-wide search
    (handles _temp-* folders parented outside the main tree).
    Returns {windows_folder_name: folder_id}.
    """
    print(f"\nResolving {len(subfolder_names)} unique subfolder names via SuiteQL...")
    name_to_id: Dict[str, int] = {}  # windows_name -> NS folder_id

    # Build a mapping of all NetSuite candidate names for each windows name
    # so we can match the query result back to the original windows key
    ns_to_win: Dict[str, str] = {}  # ns_name -> windows_name
    for win_name in subfolder_names:
        for ns_cand in _win_name_to_ns_candidates(win_name):
            ns_to_win[ns_cand] = win_name

    all_ns_names = sorted(ns_to_win.keys())

    # --- Pass 1: search under ROOT_FOLDER_ID tree ---
    for i in range(0, len(all_ns_names), QUERY_BATCH_SIZE):
        batch = all_ns_names[i:i + QUERY_BATCH_SIZE]
        quoted = ', '.join(f"'{n.replace(chr(39), chr(39)*2)}'" for n in batch)
        query = f"""
            SELECT id, name
            FROM MediaItemFolder
            WHERE parent IN (
                SELECT id FROM MediaItemFolder
                START WITH id = {ROOT_FOLDER_ID}
                CONNECT BY PRIOR id = parent
            )
            AND name IN ({quoted})
        """
        result = execute_query(query, account=account, environment=environment)
        if result.get('error'):
            print(f"  WARN: query error for batch {i//QUERY_BATCH_SIZE + 1}: {result['error']}", file=sys.stderr)
            continue
        for rec in result.get('records', []):
            ns_name = str(rec.get('name', ''))
            folder_id = int(rec.get('id', 0))
            if ns_name in ns_to_win and folder_id:
                win_name = ns_to_win[ns_name]
                name_to_id[win_name] = folder_id
        print(f"  Batch {i//QUERY_BATCH_SIZE + 1}: resolved {len(result.get('records', []))} folders")

    # --- Pass 2: broader search for anything still unresolved (temp folders, etc.) ---
    unresolved_win = [n for n in subfolder_names if n not in name_to_id]
    if unresolved_win:
        print(f"  Trying broader search for {len(unresolved_win)} unresolved folders...")
        unresolved_ns = []
        for win_name in unresolved_win:
            unresolved_ns.extend(_win_name_to_ns_candidates(win_name))
        unresolved_ns = sorted(set(unresolved_ns))

        for i in range(0, len(unresolved_ns), QUERY_BATCH_SIZE):
            batch = unresolved_ns[i:i + QUERY_BATCH_SIZE]
            quoted = ', '.join(f"'{n.replace(chr(39), chr(39)*2)}'" for n in batch)
            # No parent restriction — search entire org
            query = f"SELECT id, name FROM MediaItemFolder WHERE name IN ({quoted})"
            result = execute_query(query, account=account, environment=environment)
            if result.get('error'):
                continue
            for rec in result.get('records', []):
                ns_name = str(rec.get('name', ''))
                folder_id = int(rec.get('id', 0))
                if ns_name in ns_to_win and folder_id:
                    win_name = ns_to_win[ns_name]
                    if win_name not in name_to_id:
                        name_to_id[win_name] = folder_id
        newly = len([n for n in unresolved_win if n in name_to_id])
        print(f"  Broader search found {newly} additional folders")

    print(f"  Total resolved: {len(name_to_id)}/{len(subfolder_names)} subfolder names")
    return name_to_id


def resolve_file_ids(
    folder_ids: List[int],
    account: str,
    environment: str,
) -> Dict[Tuple[int, str], int]:
    """
    For a list of folder IDs, fetch all files and return {(folder_id, name): file_id}.
    """
    folder_to_files: Dict[Tuple[int, str], int] = {}
    folder_id_list = sorted(set(folder_ids))

    for i in range(0, len(folder_id_list), QUERY_BATCH_SIZE):
        batch = folder_id_list[i:i + QUERY_BATCH_SIZE]
        in_clause = ', '.join(str(fid) for fid in batch)
        query = f"SELECT id, name, folder FROM File WHERE folder IN ({in_clause})"
        result = execute_query(query, account=account, environment=environment)
        if result.get('error'):
            print(f"  WARN: file query error: {result['error']}", file=sys.stderr)
            continue
        for rec in result.get('records', []):
            fid = int(rec.get('folder', 0))
            fname = str(rec.get('name', ''))
            file_id = int(rec.get('id', 0))
            if fid and fname and file_id:
                folder_to_files[(fid, fname)] = file_id

    return folder_to_files


def match_entries_to_file_ids(
    entries: List[Dict],
    folder_name_to_id: Dict[str, int],
    folder_file_map: Dict[Tuple[int, str], int],
) -> None:
    """
    For each entry, find the NetSuite file ID using subfolder_name + normalized filename.
    First tries exact + normalized candidates; falls back to fuzzy matching within the folder.
    Sets entry['file_id'] (int) or entry['resolve_error'] (str).
    """
    # Build fuzzy lookup: folder_id -> {fuzzy_stem: (file_id, actual_name)}
    # This lets us find files even when special chars are completely different
    fuzzy_by_folder: Dict[int, Dict[str, Tuple[int, str]]] = {}
    for (fid, fname), file_id in folder_file_map.items():
        if fid not in fuzzy_by_folder:
            fuzzy_by_folder[fid] = {}
        fkey = fuzzy_normalize(fname)
        # Keep shortest name for a given fuzzy key (prefer base over dedup variants)
        if fkey not in fuzzy_by_folder[fid] or len(fname) < len(fuzzy_by_folder[fid][fkey][1]):
            fuzzy_by_folder[fid][fkey] = (file_id, fname)

    win_ext = Path('dummy').suffix  # just for type reference

    for entry in entries:
        subfolder_name = entry['subfolder_name']
        win_filename = entry['filename']

        folder_id = folder_name_to_id.get(subfolder_name)
        if folder_id is None:
            entry['file_id'] = None
            entry['resolve_error'] = f"Subfolder '{subfolder_name}' not found in NetSuite"
            continue

        entry['folder_id'] = folder_id
        candidates = normalize_filename(win_filename)
        matched_name = None
        file_id = None

        # --- Pass 1: exact + normalized candidates ---
        for candidate in candidates:
            key = (folder_id, candidate)
            if key in folder_file_map:
                file_id = folder_file_map[key]
                matched_name = candidate
                break

        # --- Pass 2: fuzzy fallback (strip all special chars, compare stems) ---
        if file_id is None and folder_id in fuzzy_by_folder:
            win_ext = Path(win_filename).suffix.lower()
            win_fuzzy = fuzzy_normalize(win_filename)
            folder_fuzzy = fuzzy_by_folder[folder_id]
            if win_fuzzy in folder_fuzzy:
                candidate_fid, candidate_name = folder_fuzzy[win_fuzzy]
                # Only accept if file extension matches
                if Path(candidate_name).suffix.lower() == win_ext:
                    file_id = candidate_fid
                    matched_name = candidate_name
                    entry['match_method'] = 'fuzzy'

        if file_id:
            entry['file_id'] = file_id
            entry['matched_name'] = matched_name
            entry['resolve_error'] = None
        else:
            entry['file_id'] = None
            entry['resolve_error'] = (
                f"File not found in folder {subfolder_name} (tried: {', '.join(candidates[:3])}{'...' if len(candidates) > 3 else ''})"
            )


# ---------------------------------------------------------------------------
# Credit card detection (regex + Luhn)
# ---------------------------------------------------------------------------

# Separator: optional space, dash, dot, or underscore (underscore appears in CC auth forms)
# Boundaries: negative digit lookahead/lookbehind instead of \b, because \b fails when the
# card number is flanked by underscores (e.g. __4259 0900 1376 1839__) since _ is a word char.
SEP = r'[-. _]?'

CC_PATTERNS = [
    ('Visa',       re.compile(r'(?<!\d)4[0-9]{3}' + SEP + r'[0-9]{4}' + SEP + r'[0-9]{4}' + SEP + r'[0-9]{4}(?!\d)')),
    ('Mastercard', re.compile(r'(?<!\d)(?:5[1-5][0-9]{2}|2[2-7][0-9]{2})' + SEP + r'[0-9]{4}' + SEP + r'[0-9]{4}' + SEP + r'[0-9]{4}(?!\d)')),
    ('Amex',       re.compile(r'(?<!\d)3[47][0-9]{2}' + SEP + r'[0-9]{6}' + SEP + r'[0-9]{5}(?!\d)')),
    ('Discover',   re.compile(r'(?<!\d)6(?:011|5[0-9]{2})' + SEP + r'[0-9]{4}' + SEP + r'[0-9]{4}' + SEP + r'[0-9]{4}(?!\d)')),
]


def luhn_check(number: str) -> bool:
    """Validate a digit string using the Luhn algorithm."""
    digits = [int(c) for c in reversed(number) if c.isdigit()]
    if len(digits) < 13:
        return False
    total = 0
    for i, d in enumerate(digits):
        if i % 2 == 1:
            d *= 2
            if d > 9:
                d -= 9
        total += d
    return total % 10 == 0


def redact_card(match_str: str) -> str:
    """Redact all but last 4 digits of a matched card number."""
    digits_only = re.sub(r'[^0-9]', '', match_str)
    sep_char = '-' if '-' in match_str else (' ' if ' ' in match_str else '')
    if len(digits_only) == 15:  # Amex
        return f"XXXX{sep_char}XXXXXX{sep_char}{digits_only[-5:]}"
    return f"XXXX{sep_char}XXXX{sep_char}XXXX{sep_char}{digits_only[-4:]}"


def scan_text_for_cards(text: str) -> List[Dict[str, str]]:
    """
    Scan text for credit card numbers. Returns list of hits with card type,
    redacted number, and surrounding context.
    """
    # Normalize Unicode whitespace to ASCII space so the separator class matches consistently.
    text = re.sub(r'[\xa0\u00a0\u2007\u202f\u2009\u200a\u2002\u2003\u2004\u2005\u2006]', ' ', text)
    # Replace underscores with spaces — CC auth forms use underscore runs as field blanks,
    # e.g. "Card Number: __4259 0900 1376 1839______". Without this, \b-free regex still
    # needs digits not adjacent to other digits, which holds, but the underscore between
    # groups (4259_0900) is included in SEP so it's already handled by the pattern.
    # We replace here anyway to normalize before context extraction.
    text = text.replace('_', ' ')
    hits = []
    seen: Set[str] = set()
    for card_type, pattern in CC_PATTERNS:
        for m in pattern.finditer(text):
            raw = m.group(0)
            digits = re.sub(r'[^0-9]', '', raw)
            if digits in seen:
                continue
            if not luhn_check(digits):
                continue
            seen.add(digits)
            start = max(0, m.start() - 30)
            end = min(len(text), m.end() + 30)
            context = text[start:end].replace('\n', ' ').strip()
            # Redact the actual number in context
            redacted = redact_card(raw)
            context_redacted = context.replace(raw, redacted)
            hits.append({
                'card_type': card_type,
                'redacted': redacted,
                'context': context_redacted,
            })
    return hits


# ---------------------------------------------------------------------------
# Text extraction
# ---------------------------------------------------------------------------


def extract_text_pdf(file_path: Path, use_ocr: bool = True) -> Tuple[str, str]:
    """
    Extract text from PDF. Returns (text, method).
    Uses three passes in order:
    1. pdfplumber — content stream text (fast, handles most PDFs)
    2. PyPDF2 — per-page extract_text(), annotation values, AcroForm field values
       (catches text that pdfplumber misses: some annotation types, filled form fields)
    3. OCR — only if passes 1+2 yield no text at all (image-only PDFs)
    All non-OCR passes are always run and combined so no source is skipped.
    """
    texts = []
    methods = []

    # Pass 1: pdfplumber content stream
    try:
        import pdfplumber
        with pdfplumber.open(str(file_path)) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    texts.append(t)
        if texts:
            methods.append('pdfplumber')
    except Exception:
        pass

    # Pass 2: PyPDF2 — per-page text + annotations + AcroForm
    try:
        import PyPDF2
        reader = PyPDF2.PdfReader(str(file_path))
        for page in reader.pages:
            # Per-page text (different decoder than pdfplumber — catches some gaps)
            try:
                t = page.extract_text()
                if t:
                    texts.append(t)
            except Exception:
                pass
            # Annotation values (e.g. free-text annotations, widget /V values)
            try:
                if '/Annots' in page:
                    for annot_ref in page['/Annots']:
                        try:
                            obj = annot_ref.get_object()
                            for _, v in obj.items():
                                if isinstance(v, str) and v.strip():
                                    texts.append(v)
                        except Exception:
                            pass
            except Exception:
                pass
        # AcroForm field values (filled interactive form fields)
        fields = reader.get_fields() or {}
        for field in fields.values():
            v = field.get('/V', '')
            if v and str(v).strip():
                texts.append(str(v).strip())
        if fields:
            methods.append('acroform')
        if not methods or 'pdfplumber' not in methods:
            methods.append('pypdf2')
        else:
            methods.append('pypdf2')
    except Exception:
        pass

    combined = '\n'.join(t for t in texts if t).strip()
    if combined:
        return combined, '+'.join(methods) if methods else 'pdf'

    # Pass 3: OCR fallback for image-only PDFs
    if use_ocr:
        try:
            import pytesseract
            from pdf2image import convert_from_path
            images = convert_from_path(str(file_path), dpi=200)
            ocr_parts = [pytesseract.image_to_string(img) for img in images]
            text = '\n'.join(ocr_parts)
            if text.strip():
                return text, 'ocr'
        except Exception as e:
            return '', f'ocr_error:{e}'

    return '', 'no_text'


def extract_text_excel_xlsx(file_path: Path) -> Tuple[str, str]:
    """Extract text from Excel .xlsx by reading all cells."""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(str(file_path), read_only=True, data_only=True)
        parts = []
        for sheet in wb.worksheets:
            for row in sheet.iter_rows(values_only=True):
                for cell in row:
                    if cell is not None:
                        parts.append(str(cell))
        wb.close()
        return ' '.join(parts), 'openpyxl'
    except Exception as e:
        return '', f'openpyxl_error:{e}'


def extract_text_excel_xls(file_path: Path) -> Tuple[str, str]:
    """
    Extract text from Excel .xls (OLE2 compound file format).
    xlrd reads cell values from the worksheet stream but cannot read text in
    embedded drawings, shapes, or text boxes — those are stored in separate
    OLE streams. Raw binary string extraction from the compound file catches
    all ASCII and UTF-16LE text regardless of where it is stored.
    Both passes are combined.
    """
    parts = []
    method = 'xlrd+rawstrings'
    # Pass 1: xlrd for cell values
    try:
        import xlrd
        wb = xlrd.open_workbook(str(file_path), formatting_info=False)
        for sheet in wb.sheets():
            for row_idx in range(sheet.nrows):
                for col_idx in range(sheet.ncols):
                    v = sheet.cell_value(row_idx, col_idx)
                    if v:
                        parts.append(str(v))
    except Exception:
        pass
    # Pass 2: raw binary string extraction — catches text in shapes/drawings
    try:
        with open(str(file_path), 'rb') as f:
            raw = f.read()
        parts += [s.decode('ascii', 'ignore') for s in re.findall(rb'[\x20-\x7e]{4,}', raw)]
        parts += [s.decode('utf-16-le', 'ignore') for s in re.findall(rb'(?:[\x20-\x7e]\x00){4,}', raw)]
    except Exception:
        pass
    return '\n'.join(parts), method


def extract_text_word(file_path: Path) -> Tuple[str, str]:
    """
    Extract text from Word .docx.
    python-docx reads paragraphs and table cells but misses text boxes, headers,
    footers, and SDT content controls. Raw XML extraction via zipfile catches those.
    Both passes are combined so neither source is skipped.
    """
    import zipfile as _zipfile
    parts = []
    try:
        import docx
        doc = docx.Document(str(file_path))
        for para in doc.paragraphs:
            if para.text:
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        parts.append(cell.text)
    except Exception:
        pass
    # Raw XML pass — catches text boxes, headers/footers, SDT, and content controls
    # that python-docx does not expose through its object model.
    try:
        with _zipfile.ZipFile(str(file_path), 'r') as z:
            for name in z.namelist():
                if name.endswith('.xml'):
                    try:
                        raw = z.read(name).decode('utf-8', 'ignore')
                        parts.extend(re.findall(r'<w:t[^>]*>([^<]*)</w:t>', raw))
                    except Exception:
                        pass
    except Exception:
        pass
    method = 'docx+rawxml' if parts else 'docx_empty'
    return '\n'.join(parts), method


def extract_text_msg(file_path: Path) -> Tuple[str, str]:
    """Extract text from Outlook .msg."""
    try:
        import extract_msg
        msg = extract_msg.openMsg(str(file_path))
        parts = [
            msg.subject or '',
            msg.body or '',
            msg.htmlBody.decode('utf-8', errors='replace') if msg.htmlBody else '',
        ]
        return '\n'.join(parts), 'extract-msg'
    except Exception as e:
        return '', f'msg_error:{e}'


def extract_text(file_path: Path, use_ocr: bool = True) -> Tuple[str, str]:
    """Dispatch to the right extractor based on file extension."""
    ext = file_path.suffix.lower()
    if ext == '.pdf':
        return extract_text_pdf(file_path, use_ocr=use_ocr)
    elif ext == '.xlsx':
        return extract_text_excel_xlsx(file_path)
    elif ext == '.xls':
        return extract_text_excel_xls(file_path)
    elif ext in ('.docx', '.doc'):
        return extract_text_word(file_path)
    elif ext == '.msg':
        return extract_text_msg(file_path)
    elif ext == '.json':
        try:
            return file_path.read_text(errors='replace'), 'plaintext'
        except Exception as e:
            return '', f'read_error:{e}'
    else:
        # Try reading as text
        try:
            return file_path.read_text(errors='replace'), 'plaintext'
        except Exception as e:
            return '', f'read_error:{e}'


# ---------------------------------------------------------------------------
# Re-scan mode (--rescan-report): re-run extraction+scan on existing downloads
# ---------------------------------------------------------------------------

def run_rescan(args: argparse.Namespace) -> None:
    """
    Load an existing scan_report.json, re-scan every entry using the improved
    extractors, and write the updated report. Files must already be downloaded
    under --download-dir. No network calls are made.

    This is the correct mode to run after fixing extraction bugs — it avoids a
    full re-download of 500+ files and lets you diff CONFIRMED/NOT_FOUND counts.
    """
    report_path = Path(args.rescan_report)
    download_dir = Path(args.download_dir)
    output_dir = Path(args.output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    print(f"\n=== Re-scan mode: loading {report_path} ===")
    with open(report_path) as f:
        old_report = json.load(f)

    # Flatten all entries from every category
    all_entries: List[Dict] = []
    for category in ('confirmed', 'not_found', 'errors', 'unresolved'):
        for r in old_report.get(category, []):
            entry = dict(r)
            # Restore local_path from download_dir if not set
            if not entry.get('local_path'):
                p = download_dir / entry['subfolder_name'] / entry['filename']
                entry['local_path'] = str(p) if p.exists() else None
            all_entries.append(entry)

    print(f"  Loaded {len(all_entries)} entries "
          f"(CONFIRMED={old_report.get('summary',{}).get('CONFIRMED',0)}, "
          f"NOT_FOUND={old_report.get('summary',{}).get('NOT_FOUND',0)})")

    # Re-scan everything that has a local file
    to_scan = [e for e in all_entries if e.get('local_path') and Path(e['local_path']).exists()]
    no_file = [e for e in all_entries if not e.get('local_path') or not Path(e.get('local_path', '')).exists()]
    print(f"  Will re-scan: {len(to_scan)} files  |  No local file: {len(no_file)}")

    was_confirmed = sum(1 for e in all_entries if e.get('scan_status') == 'CONFIRMED')
    now_confirmed = 0

    for i, entry in enumerate(to_scan, 1):
        local_path = Path(entry['local_path'])
        subfolder = entry['subfolder_name']
        filename = entry['filename']
        old_status = entry.get('scan_status', '?')
        print(f"  [{i}/{len(to_scan)}] {subfolder}/{filename} [{old_status}]... ", end='', flush=True)

        text, method = extract_text(local_path, use_ocr=not args.no_ocr)
        entry['extraction_method'] = method

        if not text.strip():
            entry['scan_status'] = 'NOT_FOUND' if 'error' not in method.lower() else 'ERROR'
            entry['scan_error'] = method if 'error' in method.lower() else None
            entry['card_hits'] = []
            print(f"no text ({method})")
            continue

        hits = scan_text_for_cards(text)
        # If text was extracted but no card found, try OCR as additional pass for PDFs.
        # Some PDFs contain form template text (extractable) but the card number was
        # handwritten on a printed/faxed copy and only exists in the image layer.
        if not hits and not args.no_ocr and local_path.suffix.lower() == '.pdf':
            try:
                from pdf2image import convert_from_path
                import pytesseract
                images = convert_from_path(str(local_path), dpi=300)
                ocr_text = '\n'.join(pytesseract.image_to_string(img) for img in images)
                if ocr_text.strip():
                    ocr_hits = scan_text_for_cards(ocr_text)
                    if ocr_hits:
                        hits = ocr_hits
                        method = method + '+ocr'
                        entry['extraction_method'] = method
            except Exception:
                pass
        if hits:
            entry['scan_status'] = 'CONFIRMED'
            entry['card_hits'] = hits
            now_confirmed += 1
            marker = '' if old_status == 'CONFIRMED' else ' *** UPGRADED'
            print(f"CONFIRMED ({len(hits)} hit{'s' if len(hits)!=1 else ''}, method={method}){marker}")
        else:
            entry['scan_status'] = 'NOT_FOUND'
            entry['card_hits'] = []
            if old_status == 'CONFIRMED':
                print(f"NOT_FOUND (method={method}) *** REGRESSION - was CONFIRMED")
            else:
                print(f"not found (method={method})")

    # Entries with no local file retain their prior status
    for entry in no_file:
        if entry.get('scan_status') == 'CONFIRMED':
            now_confirmed += 1

    print(f"\n  Re-scan complete: CONFIRMED {was_confirmed} -> {now_confirmed} (+{now_confirmed - was_confirmed})")
    regressions = [e for e in to_scan if e.get('scan_status') != 'CONFIRMED'
                   and any(e.get('subfolder_name') == old.get('subfolder_name')
                           and e.get('filename') == old.get('filename')
                           for old in old_report.get('confirmed', []))]
    if regressions:
        print(f"  WARNING: {len(regressions)} regressions (previously CONFIRMED, now NOT_FOUND):")
        for r in regressions:
            print(f"    {r['subfolder_name']}/{r['filename']}")

    _write_report(all_entries, output_dir)
    _write_xlsx_report(all_entries, download_dir, output_dir)


# ---------------------------------------------------------------------------
# Main scan pipeline
# ---------------------------------------------------------------------------

def run_scan(args: argparse.Namespace) -> None:
    output_dir = Path(args.output_dir)
    download_dir = Path(args.download_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    download_dir.mkdir(parents=True, exist_ok=True)

    account = args.account
    environment = args.env

    # --- Step 1: Parse flagged list ---
    print(f"\n=== Step 1: Parsing {args.file_list} ===")
    entries = parse_flagged_list(args.file_list)
    print(f"  Parsed {len(entries)} entries")

    # --- Step 2: Resolve NetSuite IDs ---
    print(f"\n=== Step 2: Resolving NetSuite file IDs ===")
    subfolder_names = set(e['subfolder_name'] for e in entries)
    folder_name_to_id = resolve_folder_ids(subfolder_names, account, environment)

    resolved_folder_ids = list(set(folder_name_to_id.values()))
    print(f"\nFetching files in {len(resolved_folder_ids)} resolved folders...")
    folder_file_map = resolve_file_ids(resolved_folder_ids, account, environment)
    print(f"  Found {len(folder_file_map)} files across resolved folders")

    match_entries_to_file_ids(entries, folder_name_to_id, folder_file_map)

    resolved = [e for e in entries if e.get('file_id')]
    unresolved = [e for e in entries if not e.get('file_id')]
    print(f"\n  Resolved: {len(resolved)}/{len(entries)}")
    if unresolved:
        print(f"  Unresolved: {len(unresolved)} (see report for details)")

    if args.resolve_only:
        _write_report(entries, output_dir, download_only=True)
        return

    # --- Step 3: Download files ---
    if not args.skip_download:
        print(f"\n=== Step 3: Downloading {len(resolved)} files ===")
        for i, entry in enumerate(resolved, 1):
            file_id = entry['file_id']
            subfolder = entry['subfolder_name']
            filename = entry['filename']
            dest = download_dir / subfolder / filename
            dest.parent.mkdir(parents=True, exist_ok=True)

            if dest.exists() and not args.force:
                print(f"  [{i}/{len(resolved)}] {subfolder}/{filename} (cached)")
                entry['local_path'] = str(dest)
                continue

            print(f"  [{i}/{len(resolved)}] {subfolder}/{filename}... ", end='', flush=True)
            content = download_file_content(file_id, account, environment)
            if content:
                dest.write_bytes(content)
                entry['local_path'] = str(dest)
                print(f"OK ({len(content):,} bytes)")
            else:
                entry['local_path'] = None
                entry['scan_status'] = 'ERROR'
                entry['scan_error'] = 'Download failed'
                print("FAILED")
            time.sleep(RATE_LIMIT_DELAY)
    else:
        scan_dir = Path(args.scan_dir) if args.scan_dir else download_dir
        for entry in resolved:
            p = scan_dir / entry['subfolder_name'] / entry['filename']
            entry['local_path'] = str(p) if p.exists() else None
            if not p.exists():
                entry['scan_status'] = 'ERROR'
                entry['scan_error'] = f"File not found at {p}"

    # --- Step 4: Scan files ---
    print(f"\n=== Step 4: Scanning files for PCI data ===")
    to_scan = [e for e in resolved if e.get('local_path') and not e.get('scan_status')]
    for i, entry in enumerate(to_scan, 1):
        local_path = Path(entry['local_path'])
        subfolder = entry['subfolder_name']
        filename = entry['filename']
        print(f"  [{i}/{len(to_scan)}] {subfolder}/{filename}... ", end='', flush=True)

        text, method = extract_text(local_path, use_ocr=not args.no_ocr)
        if not text.strip():
            entry['scan_status'] = 'NOT_FOUND' if 'error' not in method.lower() else 'ERROR'
            entry['scan_error'] = method if 'error' in method.lower() else None
            entry['extraction_method'] = method
            print(f"no text ({method})")
            continue

        hits = scan_text_for_cards(text)
        entry['extraction_method'] = method
        if hits:
            entry['scan_status'] = 'CONFIRMED'
            entry['card_hits'] = hits
            print(f"CONFIRMED ({len(hits)} hit{'s' if len(hits) != 1 else ''}, method={method})")
        else:
            entry['scan_status'] = 'NOT_FOUND'
            print(f"not found (method={method})")

    # --- Step 5: Write report ---
    _write_report(entries, output_dir)
    _write_xlsx_report(entries, download_dir, output_dir)


def _write_xlsx_report(entries: List[Dict], download_dir: Path, output_dir: Path) -> None:
    """
    Write a human-validation XLSX report with 4 sheets:
    Summary, CONFIRMED, NOT_FOUND, UNRESOLVED.
    Each data row includes a relative hyperlink to open the local file.
    """
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment
        from openpyxl.utils import get_column_letter
    except ImportError:
        print("  WARN: openpyxl not installed — skipping XLSX report", file=sys.stderr)
        return

    confirmed = [e for e in entries if e.get('scan_status') == 'CONFIRMED']
    not_found = [e for e in entries if e.get('scan_status') == 'NOT_FOUND']
    errors    = [e for e in entries if e.get('scan_status') == 'ERROR']
    unresolved = [e for e in entries if not e.get('file_id') or e.get('scan_status') == 'UNRESOLVED']

    wb = openpyxl.Workbook()

    # ---- Summary sheet ----
    ws_sum = wb.active
    ws_sum.title = 'Summary'
    ws_sum.append(['PCI File Scan Report', datetime.now().strftime('%Y-%m-%d %H:%M')])
    ws_sum.append([])
    ws_sum.append(['Status', 'Count'])
    ws_sum.append(['CONFIRMED (card data found — ready for deletion)', len(confirmed)])
    ws_sum.append(['NOT_FOUND (no card data detected)', len(not_found)])
    ws_sum.append(['ERROR (could not scan — review manually)', len(errors)])
    ws_sum.append(['UNRESOLVED (could not locate in NetSuite)', len(unresolved)])
    ws_sum.append(['Total', len(entries)])

    HEADERS = ['#', 'Status', 'Subfolder', 'Filename', 'File ID',
               'Card Types', 'Hits', 'Redacted Number', 'Context', 'Open File']

    def _add_sheet(sheet_name: str, rows: List[Dict]) -> None:
        ws = wb.create_sheet(title=sheet_name)
        ws.append(HEADERS)
        for col_idx, h in enumerate(HEADERS, 1):
            ws.cell(row=1, column=col_idx).font = Font(bold=True)

        for row_num, entry in enumerate(rows, 2):
            hits = entry.get('card_hits') or []
            card_types = ', '.join(sorted(set(h.get('card_type','') for h in hits))) if hits else ''
            redacted = hits[0].get('redacted', '') if hits else ''
            context = hits[0].get('context', '') if hits else ''

            row_data = [
                row_num - 1,
                entry.get('scan_status', 'UNRESOLVED'),
                entry.get('subfolder_name', ''),
                entry.get('filename', ''),
                entry.get('file_id'),
                card_types,
                len(hits),
                redacted,
                context[:120] if context else '',
                '📂 Open',
            ]
            ws.append(row_data)

            # Hyperlink to local file — use path relative to the XLSX file location
            # so the link survives being moved as long as the folder structure stays intact.
            local_path = entry.get('local_path')
            if local_path and Path(local_path).exists():
                try:
                    rel = Path(local_path).resolve().relative_to(download_dir.resolve())
                    link_path = str(rel)  # e.g. "10045155/twistedxcreditcard.xls.xlsx"
                except ValueError:
                    link_path = Path(local_path).resolve().as_uri()
                ws.cell(row=row_num, column=10).hyperlink = link_path
                ws.cell(row=row_num, column=10).font = Font(color='0563C1', underline='single')

        # Column widths
        ws.column_dimensions['A'].width = 5
        ws.column_dimensions['B'].width = 12
        ws.column_dimensions['C'].width = 14
        ws.column_dimensions['D'].width = 50
        ws.column_dimensions['E'].width = 12
        ws.column_dimensions['F'].width = 14
        ws.column_dimensions['G'].width = 6
        ws.column_dimensions['H'].width = 24
        ws.column_dimensions['I'].width = 60
        ws.column_dimensions['J'].width = 10

    _add_sheet(f'CONFIRMED ({len(confirmed)})', confirmed)
    _add_sheet(f'NOT_FOUND ({len(not_found)})', not_found)
    if unresolved:
        _add_sheet(f'UNRESOLVED ({len(unresolved)})', unresolved)

    xlsx_path = download_dir / 'pci_scan_report.xlsx'
    wb.save(str(xlsx_path))
    print(f"  XLSX:    {xlsx_path}")


def _write_report(entries: List[Dict], output_dir: Path, download_only: bool = False) -> None:
    """Write JSON report and human-readable summary."""
    now = datetime.now().isoformat()

    # Categorize
    confirmed = [e for e in entries if e.get('scan_status') == 'CONFIRMED']
    not_found = [e for e in entries if e.get('scan_status') == 'NOT_FOUND']
    errors = [e for e in entries if e.get('scan_status') == 'ERROR']
    unresolved = [e for e in entries if not e.get('file_id')]

    report = {
        'generated_at': now,
        'total_entries': len(entries),
        'summary': {
            'CONFIRMED': len(confirmed),
            'NOT_FOUND': len(not_found),
            'ERROR': len(errors),
            'UNRESOLVED': len(unresolved),
        },
        'confirmed': [_entry_to_report(e) for e in confirmed],
        'not_found': [_entry_to_report(e) for e in not_found],
        'errors': [_entry_to_report(e) for e in errors],
        'unresolved': [_entry_to_report(e) for e in unresolved],
    }

    report_path = output_dir / 'scan_report.json'
    with open(report_path, 'w') as f:
        json.dump(report, f, indent=2)

    # Human-readable summary
    summary_path = output_dir / 'scan_summary.txt'
    with open(summary_path, 'w') as f:
        f.write(f"PCI Scan Report — {now}\n")
        f.write("=" * 60 + "\n\n")
        f.write(f"Total entries:   {len(entries)}\n")
        f.write(f"CONFIRMED:       {len(confirmed)} (card data found — ready for deletion)\n")
        f.write(f"NOT_FOUND:       {len(not_found)} (no card data detected)\n")
        f.write(f"ERROR:           {len(errors)} (could not scan — review manually)\n")
        f.write(f"UNRESOLVED:      {len(unresolved)} (could not locate in NetSuite)\n\n")

        if confirmed:
            f.write("--- CONFIRMED FILES (delete these) ---\n")
            for e in confirmed:
                hits = e.get('card_hits', [])
                f.write(f"  file_id={e.get('file_id')} {e['subfolder_name']}/{e['filename']}\n")
                for h in hits[:3]:
                    f.write(f"    [{h['card_type']}] {h['redacted']} — \"{h['context'][:80]}\"\n")
            f.write("\n")

        if unresolved:
            f.write("--- UNRESOLVED (check manually) ---\n")
            for e in unresolved:
                f.write(f"  {e['subfolder_name']}/{e['filename']} — {e.get('resolve_error', 'unknown')}\n")
            f.write("\n")

        if errors:
            f.write("--- ERRORS (review manually) ---\n")
            for e in errors:
                f.write(f"  {e['subfolder_name']}/{e['filename']} — {e.get('scan_error', 'unknown')}\n")
            f.write("\n")

    print(f"\n=== Report written ===")
    print(f"  JSON:    {report_path}")
    print(f"  Summary: {summary_path}")
    print(f"\n  CONFIRMED (ready to delete): {len(confirmed)}")
    print(f"  NOT_FOUND:                   {len(not_found)}")
    print(f"  ERROR:                       {len(errors)}")
    print(f"  UNRESOLVED:                  {len(unresolved)}")
    if confirmed:
        print(f"\n  Run delete_file.py --report {report_path} --dry-run to preview deletions.")


def _entry_to_report(entry: Dict) -> Dict:
    """Serialize a flagged entry for the JSON report."""
    return {
        'subfolder_name': entry.get('subfolder_name', ''),
        'filename': entry.get('filename', ''),
        'file_id': entry.get('file_id'),
        'folder_id': entry.get('folder_id'),
        'matched_name': entry.get('matched_name'),
        'scan_status': entry.get('scan_status', 'UNRESOLVED'),
        'scan_error': entry.get('scan_error'),
        'resolve_error': entry.get('resolve_error'),
        'extraction_method': entry.get('extraction_method'),
        'card_hits': entry.get('card_hits', []),
        'local_path': entry.get('local_path'),
    }


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main() -> None:
    parser = argparse.ArgumentParser(
        description='Scan NetSuite Filing Cabinet files for PCI credit card data.'
    )
    parser.add_argument('--file-list', default=None,
                        help='Path to pipe-delimited flagged file list (e.g. /tmp/confirmed_files.txt). '
                             'Required unless --rescan-report is set.')
    parser.add_argument('--output-dir', default='/tmp/pci-scan-results',
                        help='Directory for scan report and summary (default: /tmp/pci-scan-results)')
    parser.add_argument('--download-dir', default='/tmp/pci-scan-downloads',
                        help='Directory to download files into (default: /tmp/pci-scan-downloads). '
                             'WARNING: downloaded files may contain PCI-sensitive data. '
                             'Use an encrypted volume in shared/production environments.')
    parser.add_argument('--account', default=DEFAULT_ACCOUNT,
                        help=f'NetSuite account (default: {DEFAULT_ACCOUNT})')
    parser.add_argument('--env', default=DEFAULT_ENVIRONMENT,
                        help=f'NetSuite environment (default: {DEFAULT_ENVIRONMENT})')
    parser.add_argument('--resolve-only', action='store_true',
                        help='Only resolve file IDs, do not download or scan')
    parser.add_argument('--skip-download', action='store_true',
                        help='Skip download step, scan files already in --scan-dir')
    parser.add_argument('--scan-dir', default=None,
                        help='Directory of pre-downloaded files (used with --skip-download)')
    parser.add_argument('--no-ocr', action='store_true',
                        help='Disable OCR fallback for scanned PDFs')
    parser.add_argument('--force', action='store_true',
                        help='Re-download files even if they already exist locally')
    parser.add_argument('--rescan-report', default=None, metavar='REPORT_JSON',
                        help='Re-scan existing downloads using an existing scan_report.json. '
                             'No network calls; only re-runs text extraction and card detection. '
                             'Use after fixing extraction bugs to avoid a full re-download.')
    args = parser.parse_args()

    print(f"NetSuite PCI File Scanner")
    print(f"  Output dir:   {args.output_dir}")
    print(f"  Download dir: {args.download_dir}")
    if args.download_dir.startswith('/tmp'):
        print(f"  WARNING: --download-dir is under /tmp. Downloaded files may contain PCI card data.")
        print(f"           Use an encrypted volume on shared/production systems.")

    if args.rescan_report:
        print(f"  Mode:         re-scan ({args.rescan_report})")
        run_rescan(args)
    else:
        print(f"  File list:    {args.file_list}")
        print(f"  Account:      {args.account} / {args.env}")
        run_scan(args)


if __name__ == '__main__':
    main()
